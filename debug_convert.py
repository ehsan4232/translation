"""
تبدیل‌کننده ساده و مؤثر MD به DOCX
بدون پیچیدگی - فقط کار کنه!
"""

import os
import glob
from pathlib import Path
import re

def install_packages():
    try:
        import docx
        return True
    except ImportError:
        print("نصب python-docx...")
        os.system("pip install python-docx")
        return False

def simple_md_to_docx():
    if not install_packages():
        print("دوباره اجرا کنید.")
        return
    
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    # پیدا کردن فایل‌ها
    md_files = glob.glob("*.md")
    if not md_files:
        print("❌ فایل MD پیدا نشد!")
        return
    
    # ایجاد پوشه خروجی
    output_dir = Path("docx_output")
    output_dir.mkdir(exist_ok=True)
    
    print(f"📚 {len(md_files)} فایل پیدا شد...")
    
    for md_file in md_files:
        print(f"🔄 {md_file}")
        
        try:
            # خواندن فایل
            with open(md_file, 'r', encoding='utf-8') as f:
                content = f.read()
            
            # ایجاد سند جدید
            doc = Document()
            
            # پاک کردن HTML tags
            content = re.sub(r'<[^>]+>', '', content)
            
            # تقسیم به خطوط
            lines = content.split('\n')
            
            for line in lines:
                original_line = line
                line = line.strip()
                
                print(f"پردازش: '{line[:50]}...'")  # Debug
                
                # خط خالی
                if not line:
                    doc.add_paragraph()
                    print("  -> خط خالی اضافه شد")
                    continue
                
                # عنوان H1
                if line.startswith('# '):
                    title = line[2:].strip()
                    h = doc.add_heading(title, level=1)
                    # تنظیم فونت
                    run = h.runs[0]
                    run.font.name = 'Vazir'
                    run.font.size = Pt(20)
                    run.font.color.rgb = RGBColor(0, 51, 102)
                    h.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    print(f"  -> H1: {title}")
                    continue
                
                # عنوان H2
                elif line.startswith('## '):
                    title = line[3:].strip()
                    h = doc.add_heading(title, level=2)
                    # تنظیم فونت
                    run = h.runs[0]
                    run.font.name = 'Vazir'
                    run.font.size = Pt(18)
                    run.font.color.rgb = RGBColor(51, 102, 153)
                    h.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    print(f"  -> H2: {title}")
                    continue
                
                # عنوان H3
                elif line.startswith('### '):
                    title = line[4:].strip()
                    h = doc.add_heading(title, level=3)
                    # تنظیم فونت
                    run = h.runs[0]
                    run.font.name = 'Vazir'
                    run.font.size = Pt(16)
                    run.font.color.rgb = RGBColor(102, 102, 102)
                    h.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    print(f"  -> H3: {title}")
                    continue
                
                # عنوان H4
                elif line.startswith('#### '):
                    title = line[5:].strip()
                    h = doc.add_heading(title, level=4)
                    # تنظیم فونت
                    run = h.runs[0]
                    run.font.name = 'Vazir'
                    run.font.size = Pt(14)
                    run.font.color.rgb = RGBColor(136, 136, 136)
                    h.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    print(f"  -> H4: {title}")
                    continue
                
                # نقل قول
                elif line.startswith('>'):
                    quote_text = line[1:].strip()
                    p = doc.add_paragraph(quote_text)
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    # فونت
                    for run in p.runs:
                        run.font.name = 'Vazir'
                        run.font.size = Pt(11)
                        run.font.italic = True
                        run.font.color.rgb = RGBColor(85, 85, 85)
                    print(f"  -> Quote: {quote_text}")
                    continue
                
                # لیست
                elif line.startswith('- ') or line.startswith('* '):
                    list_text = line[2:].strip()
                    p = doc.add_paragraph(f"• {list_text}")
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    for run in p.runs:
                        run.font.name = 'Vazir'
                        run.font.size = Pt(12)
                    print(f"  -> List: {list_text}")
                    continue
                
                # لیست شماره‌دار
                elif re.match(r'^\d+\. ', line):
                    list_text = re.sub(r'^\d+\. ', '', line).strip()
                    p = doc.add_paragraph(list_text)
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    for run in p.runs:
                        run.font.name = 'Vazir'
                        run.font.size = Pt(12)
                    print(f"  -> Numbered: {list_text}")
                    continue
                
                # خط جداکننده
                elif line.startswith('---') or line.startswith('***'):
                    p = doc.add_paragraph('─' * 50)
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    print("  -> Separator")
                    continue
                
                # متن عادی
                else:
                    p = doc.add_paragraph(line)
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    
                    # تنظیم فونت برای همه runs
                    if not p.runs:
                        p.add_run(line)
                    
                    for run in p.runs:
                        run.font.name = 'Vazir'
                        run.font.size = Pt(12)
                        run.font.color.rgb = RGBColor(0, 0, 0)
                    
                    print(f"  -> Normal: {line[:30]}...")
            
            # ذخیره
            output_file = output_dir / f"{Path(md_file).stem}.docx"
            doc.save(str(output_file))
            print(f"✅ ذخیره شد: {output_file}")
            
        except Exception as e:
            print(f"❌ خطا: {e}")
            import traceback
            traceback.print_exc()

if __name__ == "__main__":
    print("🚀 تبدیل ساده MD به DOCX")
    simple_md_to_docx()
