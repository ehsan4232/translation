"""
تبدیل‌کننده نهایی MD به DOCX - تضمین شده!
"""

import os
import glob
from pathlib import Path
import re

def ensure_packages():
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.style import WD_STYLE_TYPE
        return True
    except ImportError:
        print("نصب python-docx...")
        os.system("pip install python-docx")
        print("دوباره اجرا کنید.")
        return False

def convert_md_files():
    if not ensure_packages():
        return
    
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    
    # پیدا کردن فایل‌ها
    md_files = glob.glob("*.md")
    if not md_files:
        print("❌ فایل MD پیدا نشد!")
        return
    
    output_dir = Path("docx_output")
    output_dir.mkdir(exist_ok=True)
    
    print(f"📚 {len(md_files)} فایل پیدا شد...")
    
    for md_file in md_files:
        print(f"\n🔄 تبدیل: {md_file}")
        
        try:
            # خواندن فایل
            with open(md_file, 'r', encoding='utf-8') as f:
                content = f.read()
            
            # سند جدید
            doc = Document()
            
            # حذف HTML tags
            clean_content = re.sub(r'<[^>]+>', '', content)
            
            # تقسیم به خطوط
            lines = clean_content.split('\n')
            
            line_count = 0
            for line in lines:
                line_count += 1
                original_line = line
                line = line.strip()
                
                # خط خالی = فاصله
                if not line:
                    doc.add_paragraph("")
                    continue
                
                # H1
                if line.startswith('# '):
                    text = line[2:].strip()
                    para = doc.add_paragraph()
                    run = para.add_run(text)
                    run.font.name = 'Vazir'
                    run.font.size = Pt(20)
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(0, 51, 102)
                    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    print(f"   H1: {text}")
                
                # H2
                elif line.startswith('## '):
                    text = line[3:].strip()
                    para = doc.add_paragraph()
                    run = para.add_run(text)
                    run.font.name = 'Vazir'
                    run.font.size = Pt(18)
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(51, 102, 153)
                    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    print(f"   H2: {text}")
                
                # H3
                elif line.startswith('### '):
                    text = line[4:].strip()
                    para = doc.add_paragraph()
                    run = para.add_run(text)
                    run.font.name = 'Vazir'
                    run.font.size = Pt(16)
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(102, 102, 102)
                    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    print(f"   H3: {text}")
                
                # H4
                elif line.startswith('#### '):
                    text = line[5:].strip()
                    para = doc.add_paragraph()
                    run = para.add_run(text)
                    run.font.name = 'Vazir'
                    run.font.size = Pt(14)
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(136, 136, 136)
                    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    print(f"   H4: {text}")
                
                # نقل قول
                elif line.startswith('> '):
                    text = line[2:].strip()
                    para = doc.add_paragraph()
                    run = para.add_run(f"» {text}")
                    run.font.name = 'Vazir'
                    run.font.size = Pt(11)
                    run.font.italic = True
                    run.font.color.rgb = RGBColor(85, 85, 85)
                    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    para.left_indent = Pt(36)
                    print(f"   Quote: {text}")
                
                # لیست نقطه‌ای
                elif line.startswith('- ') or line.startswith('* '):
                    text = line[2:].strip()
                    para = doc.add_paragraph()
                    run = para.add_run(f"• {text}")
                    run.font.name = 'Vazir'
                    run.font.size = Pt(12)
                    run.font.color.rgb = RGBColor(0, 0, 0)
                    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    para.left_indent = Pt(18)
                    print(f"   List: {text}")
                
                # لیست شماره‌دار
                elif re.match(r'^\d+\. ', line):
                    text = re.sub(r'^\d+\. ', '', line).strip()
                    para = doc.add_paragraph()
                    run = para.add_run(text)
                    run.font.name = 'Vazir'
                    run.font.size = Pt(12)
                    run.font.color.rgb = RGBColor(0, 0, 0)
                    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    para.left_indent = Pt(18)
                    print(f"   Numbered: {text}")
                
                # خط جداکننده
                elif line.startswith('---') or line.startswith('***'):
                    para = doc.add_paragraph()
                    run = para.add_run('─' * 40)
                    run.font.name = 'Vazir'
                    run.font.size = Pt(12)
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    print("   Separator")
                
                # متن عادی
                else:
                    # پردازش bold و italic
                    para = doc.add_paragraph()
                    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    
                    # پیدا کردن **bold** و *italic*
                    parts = re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)', line)
                    
                    for part in parts:
                        if part.startswith('**') and part.endswith('**'):
                            # Bold
                            text = part[2:-2]
                            run = para.add_run(text)
                            run.font.name = 'Vazir'
                            run.font.size = Pt(12)
                            run.font.bold = True
                            run.font.color.rgb = RGBColor(0, 0, 0)
                        elif part.startswith('*') and part.endswith('*'):
                            # Italic
                            text = part[1:-1]
                            run = para.add_run(text)
                            run.font.name = 'Vazir'
                            run.font.size = Pt(12)
                            run.font.italic = True
                            run.font.color.rgb = RGBColor(0, 0, 0)
                        elif part.startswith('`') and part.endswith('`'):
                            # Code
                            text = part[1:-1]
                            run = para.add_run(text)
                            run.font.name = 'Consolas'
                            run.font.size = Pt(10)
                            run.font.color.rgb = RGBColor(139, 69, 19)
                        else:
                            # Normal text
                            if part.strip():
                                run = para.add_run(part)
                                run.font.name = 'Vazir'
                                run.font.size = Pt(12)
                                run.font.color.rgb = RGBColor(0, 0, 0)
                    
                    if line:
                        print(f"   Normal: {line[:40]}...")
            
            # ذخیره
            output_file = output_dir / f"{Path(md_file).stem}.docx"
            doc.save(str(output_file))
            print(f"✅ ذخیره شد: {output_file}")
            print(f"   تعداد خطوط پردازش شده: {line_count}")
            
        except Exception as e:
            print(f"❌ خطا در {md_file}: {e}")
            import traceback
            traceback.print_exc()
    
    print(f"\n🎉 تبدیل کامل!")
    print(f"📁 فایل‌ها در پوشه docx_output ذخیره شدند.")
    print(f"\n📝 فونت‌های استفاده شده:")
    print(f"   • Vazir برای همه متن‌ها و عناوین")
    print(f"   • H1: 20pt Bold آبی تیره")
    print(f"   • H2: 18pt Bold آبی متوسط")
    print(f"   • H3: 16pt Bold خاکستری")
    print(f"   • H4: 14pt Bold خاکستری روشن")
    print(f"   • Normal: 12pt مشکی")

if __name__ == "__main__":
    print("🚀 تبدیل MD به DOCX - نسخه نهایی")
    print("=" * 50)
    convert_md_files()
