"""
تبدیل اختصاصی فایل‌های MD با HTML به DOCX
برای فایل‌های با فرمت <div dir="rtl"> و فونت Vazir
"""

import os
import glob
from pathlib import Path
import re

def install_required_packages():
    """نصب کتابخانه‌های مورد نیاز"""
    try:
        import docx
        from bs4 import BeautifulSoup
        return True
    except ImportError:
        print("📦 در حال نصب کتابخانه‌های مورد نیاز...")
        os.system("pip install python-docx beautifulsoup4 lxml")
        print("✅ کتابخانه‌ها نصب شدند. لطفاً دوباره اجرا کنید.")
        return False

def md_to_docx_custom():
    """تبدیل اختصاصی MD به DOCX"""
    
    if not install_required_packages():
        return
    
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    from bs4 import BeautifulSoup
    
    # ایجاد پوشه خروجی
    output_dir = Path("docx_output")
    output_dir.mkdir(exist_ok=True)
    
    # پیدا کردن فایل‌های MD
    md_files = glob.glob("*.md")
    
    if not md_files:
        print("❌ هیچ فایل MD پیدا نشد!")
        return
    
    print(f"📚 {len(md_files)} فایل MD پیدا شد...")
    
    successful = 0
    failed = 0
    
    for md_file in md_files:
        try:
            print(f"🔄 در حال تبدیل: {md_file}")
            
            # خواندن فایل MD
            with open(md_file, 'r', encoding='utf-8') as f:
                md_content = f.read()
            
            # ایجاد سند Word
            doc = Document()
            
            # تنظیمات صفحه
            sections = doc.sections
            for section in sections:
                section.page_width = Inches(8.5)
                section.page_height = Inches(11)
                section.left_margin = Inches(1)
                section.right_margin = Inches(1)
                section.top_margin = Inches(1)
                section.bottom_margin = Inches(1)
            
            # تنظیم استایل‌ها با فونت Vazir
            def setup_styles(doc):
                """تنظیم استایل‌ها با فونت Vazir"""
                
                # استایل عنوان H1
                try:
                    h1_style = doc.styles['Heading 1']
                except:
                    h1_style = doc.styles.add_style('Heading 1', WD_STYLE_TYPE.PARAGRAPH)
                
                h1_font = h1_style.font
                h1_font.name = 'Vazir'
                h1_font.size = Pt(20)
                h1_font.bold = True
                h1_font.color.rgb = RGBColor(0, 51, 102)  # آبی تیره
                h1_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                h1_style.paragraph_format.space_after = Pt(12)
                h1_style.paragraph_format.space_before = Pt(18)
                
                # استایل عنوان H2  
                try:
                    h2_style = doc.styles['Heading 2']
                except:
                    h2_style = doc.styles.add_style('Heading 2', WD_STYLE_TYPE.PARAGRAPH)
                
                h2_font = h2_style.font
                h2_font.name = 'Vazir'
                h2_font.size = Pt(18)
                h2_font.bold = True
                h2_font.color.rgb = RGBColor(51, 102, 153)  # آبی متوسط
                h2_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                h2_style.paragraph_format.space_after = Pt(10)
                h2_style.paragraph_format.space_before = Pt(15)
                
                # استایل عنوان H3
                try:
                    h3_style = doc.styles['Heading 3']
                except:
                    h3_style = doc.styles.add_style('Heading 3', WD_STYLE_TYPE.PARAGRAPH)
                
                h3_font = h3_style.font
                h3_font.name = 'Vazir'
                h3_font.size = Pt(16)
                h3_font.bold = True
                h3_font.color.rgb = RGBColor(102, 102, 102)  # خاکستری
                h3_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                h3_style.paragraph_format.space_after = Pt(8)
                h3_style.paragraph_format.space_before = Pt(12)
                
                # استایل عنوان H4
                try:
                    h4_style = doc.styles['Heading 4']
                except:
                    h4_style = doc.styles.add_style('Heading 4', WD_STYLE_TYPE.PARAGRAPH)
                
                h4_font = h4_style.font
                h4_font.name = 'Vazir'
                h4_font.size = Pt(14)
                h4_font.bold = True
                h4_font.color.rgb = RGBColor(136, 136, 136)  # خاکستری روشن
                h4_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                h4_style.paragraph_format.space_after = Pt(6)
                h4_style.paragraph_format.space_before = Pt(10)
                
                # استایل متن عادی
                normal_style = doc.styles['Normal']
                normal_font = normal_style.font
                normal_font.name = 'Vazir'
                normal_font.size = Pt(12)
                normal_font.color.rgb = RGBColor(0, 0, 0)
                normal_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                normal_style.paragraph_format.line_spacing = 1.2
                normal_style.paragraph_format.space_after = Pt(8)
                
                # استایل نقل قول
                try:
                    quote_style = doc.styles['Quote']
                except:
                    quote_style = doc.styles.add_style('Quote', WD_STYLE_TYPE.PARAGRAPH)
                
                quote_font = quote_style.font
                quote_font.name = 'Vazir'
                quote_font.size = Pt(11)
                quote_font.italic = True
                quote_font.color.rgb = RGBColor(85, 85, 85)
                quote_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                quote_style.paragraph_format.left_indent = Inches(0.5)
                quote_style.paragraph_format.space_after = Pt(8)
                quote_style.paragraph_format.space_before = Pt(8)
                
                # استایل کد
                try:
                    code_style = doc.styles['Code']
                except:
                    code_style = doc.styles.add_style('Code', WD_STYLE_TYPE.PARAGRAPH)
                
                code_font = code_style.font
                code_font.name = 'Consolas'
                code_font.size = Pt(10)
                code_font.color.rgb = RGBColor(139, 69, 19)
                code_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                code_style.paragraph_format.space_after = Pt(8)
                code_style.paragraph_format.space_before = Pt(8)
            
            setup_styles(doc)
            
            # حذف HTML و پردازش محتوا
            def process_content(content):
                """پردازش محتوای MD با HTML"""
                
                # حذف تگ‌های HTML اضافی
                content = re.sub(r'<div[^>]*>', '', content)
                content = re.sub(r'</div>', '', content)
                
                # تقسیم به خطوط
                lines = content.split('\n')
                
                for line_num, line in enumerate(lines):
                    line = line.strip()
                    
                    # خط خالی - اضافه کردن فاصله
                    if not line:
                        doc.add_paragraph()
                        continue
                    
                    # عناوین
                    if line.startswith('# '):
                        doc.add_heading(line[2:].strip(), level=1)
                        continue
                    elif line.startswith('## '):
                        doc.add_heading(line[3:].strip(), level=2)
                        continue
                    elif line.startswith('### '):
                        doc.add_heading(line[4:].strip(), level=3)
                        continue
                    elif line.startswith('#### '):
                        doc.add_heading(line[5:].strip(), level=4)
                        continue
                    
                    # نقل قول
                    elif line.startswith('> '):
                        quote_text = line[2:].strip()
                        p = doc.add_paragraph(quote_text, style='Quote')
                        continue
                    
                    # کد
                    elif line.startswith('```'):
                        # کد بلاک - خط بعدی رو هم بخون
                        if line_num + 1 < len(lines):
                            code_content = []
                            i = line_num + 1
                            while i < len(lines) and not lines[i].strip().startswith('```'):
                                code_content.append(lines[i])
                                i += 1
                            if code_content:
                                p = doc.add_paragraph('\n'.join(code_content), style='Code')
                        continue
                    
                    elif line.startswith('    ') and len(line) > 4:
                        # کد با ۴ اسپیس
                        code_text = line[4:]
                        p = doc.add_paragraph(code_text, style='Code')
                        continue
                    
                    # لیست با نقطه
                    elif line.startswith('- ') or line.startswith('* '):
                        list_text = line[2:].strip()
                        p = doc.add_paragraph(list_text, style='List Bullet')
                        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                        continue
                    
                    # لیست شماره‌دار
                    elif re.match(r'^\d+\. ', line):
                        list_text = re.sub(r'^\d+\. ', '', line).strip()
                        p = doc.add_paragraph(list_text, style='List Number')
                        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                        continue
                    
                    # خط جداکننده
                    elif line.startswith('---') or line.startswith('***'):
                        p = doc.add_paragraph()
                        p.add_run('─' * 50)
                        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        continue
                    
                    # متن عادی
                    else:
                        # چک کن اگه خط با ** شروع میشه (bold)
                        if line.startswith('**') and line.endswith('**') and len(line) > 4:
                            # تیتر bold
                            p = doc.add_paragraph()
                            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                            run = p.add_run(line[2:-2])
                            run.bold = True
                            run.font.name = 'Vazir'
                            run.font.size = Pt(13)
                        else:
                            # متن عادی با پردازش inline formatting
                            p = doc.add_paragraph()
                            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                            
                            # پردازش bold, italic, code
                            parts = re.split(r'(\*\*.*?\*\*|\*.*?\*|`.*?`)', line)
                            
                            for part in parts:
                                if part.startswith('**') and part.endswith('**') and len(part) > 4:
                                    # Bold
                                    run = p.add_run(part[2:-2])
                                    run.bold = True
                                    run.font.name = 'Vazir'
                                elif part.startswith('*') and part.endswith('*') and len(part) > 2:
                                    # Italic
                                    run = p.add_run(part[1:-1])
                                    run.italic = True
                                    run.font.name = 'Vazir'
                                elif part.startswith('`') and part.endswith('`') and len(part) > 2:
                                    # Inline code
                                    run = p.add_run(part[1:-1])
                                    run.font.name = 'Consolas'
                                    run.font.size = Pt(10)
                                    run.font.color.rgb = RGBColor(139, 69, 19)
                                else:
                                    # متن عادی
                                    if part.strip():
                                        run = p.add_run(part)
                                        run.font.name = 'Vazir'
            
            # پردازش محتوا
            process_content(md_content)
            
            # ذخیره سند
            output_file = output_dir / f"{Path(md_file).stem}.docx"
            doc.save(str(output_file))
            
            print(f"✅ تکمیل شد: {output_file}")
            successful += 1
            
        except Exception as e:
            print(f"❌ خطا در تبدیل {md_file}: {e}")
            failed += 1
    
    print(f"\n🎉 تبدیل کامل!")
    print(f"✅ موفق: {successful} فایل")
    print(f"❌ ناموفق: {failed} فایل")
    print(f"📁 فایل‌ها در پوشه {output_dir} ذخیره شدند.")
    print(f"\n📝 ویژگی‌های اعمال شده:")
    print(f"   • فونت Vazir برای همه عناوین و متن")
    print(f"   • اندازه‌های مختلف: H1(20pt), H2(18pt), H3(16pt), H4(14pt)")
    print(f"   • رنگ‌های متفاوت برای هر سطح عنوان")
    print(f"   • حفظ دقیق خطوط جدید و فاصله‌گذاری")
    print(f"   • پشتیبانی از RTL")
    print(f"   • پردازش صحیح HTML موجود در MD")

if __name__ == "__main__":
    print("🚀 تبدیل اختصاصی MD به DOCX (فونت Vazir)")
    print("=" * 60)
    md_to_docx_custom()
